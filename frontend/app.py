"""
Meeting Minutes Agent — Streamlit Frontend
==========================================
Communicates with the FastAPI backend (BACKEND_URL env var).
Within Container Apps the backend is on the private network (internal ingress).
"""
from __future__ import annotations

import io
import os
import re
import time
import uuid
from datetime import datetime
from typing import Optional

import requests
import streamlit as st
from azure.identity import ManagedIdentityCredential, DefaultAzureCredential
from azure.storage.blob import BlobServiceClient

# ── Configuration ─────────────────────────────────────────────────────────────
BACKEND_URL = os.environ.get("BACKEND_URL", "http://localhost:8000").rstrip("/")
AZURE_STORAGE_ACCOUNT_URL = os.environ.get("AZURE_STORAGE_ACCOUNT_URL", "")
AZURE_STORAGE_CONTAINER = os.environ.get("AZURE_STORAGE_CONTAINER", "audio-files")
POLL_INTERVAL = int(os.environ.get("POLL_INTERVAL_SECONDS", "3"))
# Max wait time for the whole pipeline. Long audio + 3 LLM stages can easily
# exceed 5 minutes, so default to 60 minutes worth of polls.
MAX_WAIT_SECONDS = int(os.environ.get("MAX_WAIT_SECONDS", "3600"))
MAX_POLLS = max(1, MAX_WAIT_SECONDS // max(1, POLL_INTERVAL))

# ── Page setup ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="会議議事録エージェント",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Custom CSS ────────────────────────────────────────────────────────────────
st.markdown(
    """
<style>
/* ── Fluent Design 2 — color tokens ──────────────────────────────────── */
:root {
    --brand-primary: #0078d4;
    --brand-hover:   #106ebe;
    --brand-pressed: #005a9e;
    --neutral-bg:    #faf9f8;
    --neutral-border:#e1dfdd;
    --neutral-fg:    #323130;
    --neutral-fg2:   #605e5c;
    --neutral-fg3:   #a19f9d;
    --success:       #107c10;
    --success-bg:    #dff6dd;
    --error:         #a4262c;
    --error-bg:      #fde7e9;
    --surface:       #ffffff;
}

/* Global font — compact */
html { font-size: 12px; }
.stApp, .stMarkdown, .stTextInput, .stTextArea, .stButton, .stTabs,
.stSelectbox, .stExpander, .stCaption, .stMetric { font-size: 0.85rem; }
h1 { font-size: 1.5rem !important; margin: 0 !important; }
h2 { font-size: 1.25rem !important; margin: 0 !important; }
h3 { font-size: 1.1rem !important; margin: 0 !important; }
h4 { font-size: 1.0rem !important; margin: 0 !important; }
p { margin-bottom: 0.3rem !important; }

/* ── Fluent buttons — all buttons get a consistent neutral style ──── */
.stButton > button,
.stDownloadButton > button {
    background: var(--surface) !important;
    color: var(--neutral-fg) !important;
    border: 1px solid var(--neutral-border) !important;
    border-radius: 4px !important;
    padding: 5px 14px !important;
    font-size: 0.82rem !important;
    font-weight: 500 !important;
    transition: background 0.15s, border-color 0.15s !important;
    box-shadow: 0 1px 2px rgba(0,0,0,0.04) !important;
}
.stButton > button:hover,
.stDownloadButton > button:hover {
    background: #f3f2f1 !important;
    border-color: #c8c6c4 !important;
}
.stButton > button:active,
.stDownloadButton > button:active {
    background: #edebe9 !important;
}
/* Primary buttons — brand blue (soft) */
.stButton > button[kind="primary"],
.stButton > button[data-testid="stBaseButton-primary"] {
    background: var(--brand-primary) !important;
    color: white !important;
    border-color: var(--brand-primary) !important;
}
.stButton > button[kind="primary"]:hover,
.stButton > button[data-testid="stBaseButton-primary"]:hover {
    background: var(--brand-hover) !important;
    border-color: var(--brand-hover) !important;
}
.stButton > button[kind="primary"]:active,
.stButton > button[data-testid="stBaseButton-primary"]:active {
    background: var(--brand-pressed) !important;
}
/* Disabled state */
.stButton > button:disabled,
.stDownloadButton > button:disabled {
    background: #f3f2f1 !important;
    color: var(--neutral-fg3) !important;
    border-color: var(--neutral-border) !important;
    box-shadow: none !important;
    cursor: not-allowed !important;
}

/* Tighten Streamlit widget gaps */
.stVerticalBlock > div { margin-bottom: 0 !important; padding-top: 0 !important; padding-bottom: 0 !important; }
.stElementContainer { margin-bottom: 0.15rem !important; }
.stTabs [data-baseweb="tab-list"] { gap: 4px; border-bottom: 2px solid var(--neutral-border); }
.stTabs [data-baseweb="tab"] { padding: 4px 10px; font-size: 0.82rem; }
.stTabs [data-baseweb="tab"][aria-selected="true"] { border-bottom: 2px solid var(--brand-primary); }
div[data-testid="stExpander"] { margin-bottom: 4px !important; }
div[data-testid="stExpander"] summary { padding: 4px 8px !important; font-size: 0.82rem; }
.stDivider { margin: 4px 0 !important; }
.stCaption { font-size: 0.78rem !important; }

/* Compact header — single line */
.app-header-compact {
    display: flex;
    align-items: center;
    gap: 12px;
    background: linear-gradient(135deg, #0078d4 0%, #004578 100%);
    color: white;
    padding: 6px 16px;
    border-radius: 6px;
    margin-bottom: 4px;
}
.app-header-compact .header-title {
    font-size: 1.15rem;
    font-weight: 700;
    white-space: nowrap;
}
.app-header-compact .header-status {
    font-size: 0.82rem;
    opacity: 0.9;
    margin-left: auto;
    white-space: nowrap;
}

/* Section headers */
.section-label {
    font-size: 0.95rem;
    font-weight: 600;
    color: var(--brand-pressed);
    margin-bottom: 2px;
}

/* Pipeline step cards */
.step-card {
    display: flex;
    align-items: center;
    gap: 8px;
    padding: 5px 10px;
    border-radius: 6px;
    border: 1px solid var(--neutral-border);
    background: var(--neutral-bg);
    margin-bottom: 4px;
    font-size: 0.78rem;
}
.step-card.done    { border-color: var(--success); background: var(--success-bg); }
.step-card.active  { border-color: var(--brand-primary); background: #deecf9; }
.step-card.error   { border-color: var(--error); background: var(--error-bg); }
.step-card.pending { border-color: var(--neutral-border); background: var(--neutral-bg); color: var(--neutral-fg3); }

/* Result markdown tweaks */
.minutes-box {
    background: var(--surface);
    border: 1px solid var(--neutral-border);
    border-radius: 6px;
    padding: 10px 14px;
}

/* Hide Streamlit branding & top bar */
#MainMenu, footer, header { visibility: hidden; height: 0; }

/* Reduce top/bottom padding */
.block-container { padding-top: 0.2rem !important; padding-bottom: 0.2rem !important; }
.stMainBlockContainer { padding-top: 0.2rem !important; }

/* Agent detail panel */
.detail-panel {
    border-left: 2px solid var(--brand-primary);
    padding-left: 8px;
}
.detail-panel h4 { color: var(--brand-primary); margin-bottom: 4px; font-size: 1.0rem !important; }

/* Topic summary line */
.topic-summary {
    color: var(--neutral-fg);
    margin: 2px 0 4px 0;
    font-size: 0.82rem;
}

/* Status panel in input page */
.status-panel {
    background: var(--neutral-bg);
    border: 1px solid var(--neutral-border);
    border-radius: 6px;
    padding: 8px;
    margin-bottom: 4px;
}
.status-ready { color: var(--success); font-weight: 600; font-size: 0.84rem; }
.status-waiting { color: var(--neutral-fg3); font-size: 0.84rem; }

/* Delete confirm button — subtle danger */
.del-confirm button {
    background: var(--error) !important;
    color: white !important;
    border-color: var(--error) !important;
}
.del-confirm button:hover {
    background: #c13530 !important;
}
</style>
""",
    unsafe_allow_html=True,
)

# ── Responsive layout via JS injection ────────────────────────────────────────
# Streamlit's st.container(height=N) only accepts px. CSS-in-markdown cannot
# reliably override inline styles in all Streamlit versions. Instead, we
# inject JS via an iframe that patches the parent document directly.
import streamlit.components.v1

streamlit.components.v1.html(
    """
    <script>
    (function() {
        var doc = window.parent.document;

        // Inject style into parent <head> (idempotent)
        if (!doc.getElementById('_vh_responsive')) {
            var s = doc.createElement('style');
            s.id = '_vh_responsive';
            s.textContent = [
                '/* Prevent page scroll */',
                '.stApp { overflow: hidden !important; height: 100vh !important; }',
                'section.stMain { overflow: hidden !important; }',
                '.stMainBlockContainer { overflow: hidden !important; }',
                '',
                '/* All Streamlit fixed-height containers → responsive */',
                'div[data-testid="stVerticalBlockBorderWrapper"] > div[style*="height"] {',
                '  height: 46vh !important;',
                '  max-height: 46vh !important;',
                '}',
                '',
                '/* Nested containers (agent detail) → smaller */',
                'div[data-testid="stVerticalBlockBorderWrapper"] > div[style*="height"]',
                '  div[data-testid="stVerticalBlockBorderWrapper"] > div[style*="height"] {',
                '  height: 38vh !important;',
                '  max-height: 38vh !important;',
                '}',
                '',
                '/* Double-nested containers (JSON previews) — keep original height */',
                'div[data-testid="stVerticalBlockBorderWrapper"] > div[style*="height"]',
                '  div[data-testid="stVerticalBlockBorderWrapper"] > div[style*="height"]',
                '    div[data-testid="stVerticalBlockBorderWrapper"] > div[style*="height"] {',
                '  /* Do not override — let the fixed px value from st.container() apply */',
                '}',
                'div[data-testid="stVerticalBlockBorderWrapper"] > div[style*="height"]',
                '  div[data-testid="stVerticalBlockBorderWrapper"] > div[style*="height"]',
                '    div[data-testid="stVerticalBlockBorderWrapper"] > div[style*="height"] {',
                '  overflow: auto !important;',
                '}',

            ].join('\\n');
            doc.head.appendChild(s);
        }

        // Belt-and-suspenders: also directly patch elements in case
        // the CSS selector doesn't match the exact DOM structure.
        function patchHeights() {
            var vh = window.parent.innerHeight;
            var mainH = Math.round(vh * 0.46);
            var subH  = Math.round(vh * 0.12);

            // Find all elements with inline height style
            doc.querySelectorAll('*').forEach(function(el) {
                var h = el.style.height;
                if (!h || !h.match(/^\\d+(\\.\\d+)?px$/)) return;
                var px = parseInt(h);
                // Skip very small (buttons etc.) and very large (full page)
                if (px < 60 || px > vh) return;
                // Tag with data attribute for identification
                if (px >= 200) {
                    // Check if already inside a main container (nested)
                    var parent = el.parentElement;
                    var depth = 0;
                    while (parent) {
                        if (parent.dataset && parent.dataset.vhPatched) depth++;
                        parent = parent.parentElement;
                    }
                    if (depth === 0) {
                        el.style.setProperty('height', mainH + 'px', 'important');
                        el.style.setProperty('max-height', mainH + 'px', 'important');
                        el.dataset.vhPatched = 'main';
                    } else if (depth === 1) {
                        var detailH = Math.round(vh * 0.38);
                        el.style.setProperty('height', detailH + 'px', 'important');
                        el.style.setProperty('max-height', detailH + 'px', 'important');
                        el.dataset.vhPatched = 'detail';
                    }
                    // depth >= 2: keep original px height (JSON preview sub-containers)
                }
            });
        }

        // Run on load, resize, and DOM mutations
        setTimeout(patchHeights, 300);
        setTimeout(patchHeights, 1000);
        window.parent.addEventListener('resize', function() {
            // Reset markers on resize so recalculation happens fresh
            doc.querySelectorAll('[data-vh-patched]').forEach(function(el) {
                delete el.dataset.vhPatched;
            });
            setTimeout(patchHeights, 100);
        });

        var observer = new MutationObserver(function() {
            setTimeout(patchHeights, 200);
        });
        observer.observe(doc.body, { childList: true, subtree: true });
    })();
    </script>
    """,
    height=0,
)

# ── Responsive vh helper ──────────────────────────────────────────────────────
# Streamlit's st.container(height=...) only accepts px. We set a rough
# px fallback for the Python call, then CSS with !important overrides
# inline styles to use viewport-relative `vh` units — fully responsive.

_SCREEN_H_FALLBACK = 900  # px — fallback for Python-side container()

def _vh(pct: int, screen_h: int = _SCREEN_H_FALLBACK) -> int:
    """Return *pct* % of screen height as px (integer)."""
    return max(80, int(screen_h * pct / 100))

# ── Session state defaults ────────────────────────────────────────────────────
# Keys here are managed manually (not bound to a widget). The toggle's
# ``show_agent_detail`` key is intentionally NOT included — Streamlit forbids
# direct assignment to a key once a widget with that key has been instantiated,
# which used to crash the "新しい議事録" reset path.
_DEFAULTS: dict = {
    "page": "input",       # input | processing | result | error
    "audio_bytes": None,
    "audio_filename": None,
    "audio_mime": None,
    "transcript_text": None,
    "input_mode": "audio",   # audio | transcript
    "job_id": None,
    "job_result": None,
    "error_msg": None,
}
for k, v in _DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ── API helpers ───────────────────────────────────────────────────────────────


@st.cache_resource
def _get_blob_service_client() -> BlobServiceClient:
    """Return a cached BlobServiceClient using Managed Identity."""
    try:
        cred = ManagedIdentityCredential()
    except Exception:
        cred = DefaultAzureCredential()
    return BlobServiceClient(account_url=AZURE_STORAGE_ACCOUNT_URL, credential=cred)


def _upload_to_blob(audio_bytes: bytes, filename: str) -> str:
    """Upload audio directly to Blob Storage and return the blob name."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "bin"
    blob_name = f"upload/{uuid.uuid4()}/{filename}"
    svc = _get_blob_service_client()
    blob_client = svc.get_blob_client(container=AZURE_STORAGE_CONTAINER, blob=blob_name)
    blob_client.upload_blob(audio_bytes, overwrite=True)
    return blob_name


def api_submit(audio_bytes: bytes, filename: str, mime: str, transcription_mode: str = "fast") -> str:
    """Upload audio to Blob Storage, then tell backend to start the pipeline."""
    blob_name = _upload_to_blob(audio_bytes, filename)
    resp = requests.post(
        f"{BACKEND_URL}/api/v1/audio/start-from-blob",
        json={
            "blob_name": blob_name,
            "filename": filename,
            "transcription_mode": transcription_mode,
        },
        timeout=30,
    )
    resp.raise_for_status()
    return resp.json()["job_id"]


def api_submit_transcript(transcript: str) -> str:
    """Submit a pre-existing transcript to backend and return job_id."""
    resp = requests.post(
        f"{BACKEND_URL}/api/v1/audio/transcript",
        json={"transcript": transcript, "language": "ja"},
        timeout=60,
    )
    resp.raise_for_status()
    return resp.json()["job_id"]


def api_list_history(limit: int = 100) -> list[dict]:
    """Return archived job entries (newest first)."""
    try:
        resp = requests.get(
            f"{BACKEND_URL}/api/v1/history",
            params={"limit": limit},
            timeout=15,
        )
        resp.raise_for_status()
        return resp.json().get("items", [])
    except requests.RequestException as exc:
        st.warning(f"履歴の取得に失敗しました: {exc}")
        return []


def api_get_history(job_id: str) -> dict | None:
    """Return the archived meta+result for *job_id*."""
    try:
        resp = requests.get(f"{BACKEND_URL}/api/v1/history/{job_id}", timeout=30)
        resp.raise_for_status()
        return resp.json()
    except requests.RequestException as exc:
        st.error(f"議事録の読み込みに失敗しました: {exc}")
        return None


def api_get_history_input(job_id: str) -> tuple[bytes, str, str] | None:
    """Return ``(bytes, filename, mime)`` of the archived input for *job_id*."""
    try:
        resp = requests.get(
            f"{BACKEND_URL}/api/v1/history/{job_id}/input", timeout=120
        )
        resp.raise_for_status()
    except requests.RequestException as exc:
        st.error(f"入力ファイルの取得に失敗しました: {exc}")
        return None
    cd = resp.headers.get("content-disposition", "")
    filename = "input.bin"
    m = re.search(r"filename\*=UTF-8''([^;]+)", cd)
    if m:
        try:
            filename = requests.utils.unquote(m.group(1))
        except Exception:
            pass
    return resp.content, filename, resp.headers.get("content-type", "application/octet-stream")


def api_delete_history(job_id: str) -> bool:
    try:
        resp = requests.delete(f"{BACKEND_URL}/api/v1/history/{job_id}", timeout=30)
        resp.raise_for_status()
        return True
    except requests.RequestException as exc:
        st.error(f"履歴の削除に失敗しました: {exc}")
        return False


def api_check_transcription_status(job_id: str) -> str | None:
    """Check Speech API status of a timed-out batch transcription."""
    try:
        resp = requests.get(f"{BACKEND_URL}/api/v1/history/{job_id}/transcription-status", timeout=30)
        resp.raise_for_status()
        return resp.json().get("speech_status", "Unknown")
    except requests.RequestException as exc:
        st.error(f"ステータス確認に失敗しました: {exc}")
        return None


def api_resume_transcription(job_id: str) -> str | None:
    """Resume a timed-out batch transcription. Returns new job_id."""
    try:
        resp = requests.post(f"{BACKEND_URL}/api/v1/history/{job_id}/resume", timeout=30)
        resp.raise_for_status()
        return resp.json().get("job_id")
    except requests.RequestException as exc:
        st.error(f"再開に失敗しました: {exc}")
        return None


# ── Transcript file parsers ───────────────────────────────────────────────────

_VTT_TIMESTAMP = re.compile(r"\d{1,2}:\d{2}(?::\d{2})?[.,]\d{3}\s*-->")
_VTT_TAG = re.compile(r"<[^>]+>")


def parse_vtt(content: str) -> str:
    """Convert WebVTT/SRT content to plain text, preserving speaker labels.

    - Strips WEBVTT header, NOTE blocks, cue identifiers, and timestamp lines.
    - Extracts speaker name from <v Speaker>...</v> tags.
    - Removes other inline tags and collapses consecutive duplicates.
    """
    lines_out: list[str] = []
    last_line: str | None = None
    for raw in content.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.upper().startswith("WEBVTT"):
            continue
        if line.startswith("NOTE"):
            continue
        if _VTT_TIMESTAMP.search(line):
            continue
        if line.isdigit():  # SRT cue numbers
            continue
        speaker = None
        m = re.match(r"<v\s+([^>]+)>\s*(.*)", line)
        if m:
            speaker = m.group(1).strip()
            line = m.group(2)
        line = _VTT_TAG.sub("", line).strip()
        if not line:
            continue
        if speaker:
            line = f"{speaker}：{line}"
        if line == last_line:
            continue
        lines_out.append(line)
        last_line = line
    return "\n".join(lines_out)


def parse_docx(data: bytes) -> str:
    """Extract plain text from a .docx file (paragraphs + table cells)."""
    from docx import Document  # lazy import

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def api_poll(job_id: str) -> dict:
    resp = requests.get(f"{BACKEND_URL}/api/v1/audio/jobs/{job_id}", timeout=30)
    resp.raise_for_status()
    return resp.json()


# ── Shared header ─────────────────────────────────────────────────────────────

def render_header(status: str = "") -> None:
    """Compact single-line header with optional status on the right."""
    status_html = (
        f'<span class="header-status">{status}</span>' if status else ""
    )
    st.markdown(
        f'<div class="app-header-compact">'
        f'<span class="header-title">会議議事録エージェント</span>'
        f'{status_html}'
        f'</div>',
        unsafe_allow_html=True,
    )


# ── Agent detail panel ─────────────────────────────────────────────────────────

def render_agent_detail(result: dict, input_mode: str = "audio", height: int | None = None) -> None:
    """Render agent input/output detail panel in the right column."""
    if height is None:
        height = _vh(40)
    cu = result.get("content_analysis")
    script = result.get("script")
    minutes = result.get("minutes")
    final = result.get("final_minutes")

    st.markdown('<div class="detail-panel">', unsafe_allow_html=True)
    st.markdown("#### エージェント入出力")

    # Make the panel itself independently scrollable so it doesn't push the
    # main column down. ``st.container(height=...)`` clips overflow and shows
    # a vertical scrollbar local to this panel only.
    with st.container(height=height, border=False):
        # Step 1: Speech Transcription — only shown when audio was provided.
        if input_mode != "transcript":
            with st.expander(
                "Step 1: 音声解析",
                expanded=cu is not None and script is None,
            ):
                st.markdown("**入力**")
                st.caption("音声ファイル（バイナリデータ）")
                st.markdown("**出力**")
                if cu:
                    with st.container(height=160):
                        st.json(cu)
                else:
                    st.info("処理待ち")

        # Step 2: Script generation
        with st.expander(
            "Step 2: スクリプト生成",
            expanded=script is not None and minutes is None,
        ):
            if input_mode == "transcript":
                st.markdown("**入力** — 文字起こしテキスト")
            else:
                st.markdown("**入力** — 音声解析結果")
            if cu:
                with st.container(height=160):
                    st.json(cu)
            else:
                st.info("入力待ち" if input_mode == "transcript" else "Step 1 の完了待ち")
            st.markdown("**出力**")
            if script:
                with st.container(height=160):
                    st.json(script)
            else:
                st.info("処理待ち")

        # Step 3: Minutes creation
        with st.expander(
            "Step 3: 議事録作成",
            expanded=minutes is not None and final is None,
        ):
            st.markdown("**入力** — スクリプト")
            if script:
                with st.container(height=160):
                    st.json(script)
            else:
                st.info("Step 2 の完了待ち")
            st.markdown("**出力**")
            if minutes:
                with st.container(height=160):
                    st.json(minutes)
            else:
                st.info("処理待ち")

        # Step 4: Terminology enrichment
        with st.expander("Step 4: 用語補足", expanded=final is not None):
            st.markdown("**入力** — 議事録")
            if minutes:
                with st.container(height=160):
                    st.json(minutes)
            else:
                st.info("Step 3 の完了待ち")
            st.markdown("**出力**")
            if final:
                with st.container(height=160):
                    st.json(final)
            else:
                st.info("処理待ち")

    st.markdown('</div>', unsafe_allow_html=True)


# ── Pages ─────────────────────────────────────────────────────────────────────

def page_input() -> None:
    render_header()

    audio_bytes: Optional[bytes] = None
    filename: Optional[str] = None
    mime: Optional[str] = None
    transcript_text: Optional[str] = None

    col_op, col_status = st.columns(2)

    # ── Left: Input operations ────────────────────────────────────────────────
    with col_op:
        with st.container(height=_vh(46), border=False):
            st.markdown('<p class="section-label">入力</p>', unsafe_allow_html=True)
            tab_record, tab_upload, tab_text = st.tabs(
                ["録音", "アップロード", "テキスト"]
            )

            with tab_record:
                st.caption("ブラウザのマイクで録音します")
                try:
                    recorded = st.audio_input("クリックして録音を開始")
                    if recorded is not None:
                        audio_bytes = recorded.read()
                        filename = "recording.wav"
                        mime = recorded.type or "audio/wav"
                        st.audio(audio_bytes, format=mime)
                        st.success(f"録音完了：{len(audio_bytes)/1024:.1f} KB")
                except AttributeError:
                    st.warning("このバージョンは audio_input 非対応です。アップロードタブをお使いください。")

            with tab_upload:
                st.caption("WAV, MP3, MP4, M4A, OGG, WebM, FLAC（最大 500 MB）")
                uploaded = st.file_uploader(
                    "ファイルを選択またはドロップ",
                    type=["wav", "mp3", "mp4", "m4a", "ogg", "webm", "flac"],
                    accept_multiple_files=False,
                )
                if uploaded is not None:
                    audio_bytes = uploaded.read()
                    filename = uploaded.name
                    mime = uploaded.type or "audio/wav"
                    st.info(f"**{filename}**（{len(audio_bytes)/1024/1024:.1f} MB）")

            with tab_text:
                st.caption("文字起こし済みテキストを貼り付け。話者：発言内容 の形式推奨")
                text_input = st.text_area(
                    "文字起こしテキスト",
                    height=100,
                    placeholder="話者１：本日はお集まりいただきまして...",
                )
                uploaded_text = st.file_uploader(
                    "テキストファイル (.txt/.md/.vtt/.srt/.docx)",
                    type=["txt", "md", "vtt", "srt", "docx"],
                    accept_multiple_files=False,
                    key="transcript_file",
                )
                if uploaded_text is not None:
                    ext = uploaded_text.name.rsplit(".", 1)[-1].lower() if "." in uploaded_text.name else ""
                    try:
                        raw_bytes = uploaded_text.read()
                        if ext == "docx":
                            text_input = parse_docx(raw_bytes)
                        else:
                            decoded = raw_bytes.decode("utf-8", errors="replace")
                            if ext in ("vtt", "srt"):
                                text_input = parse_vtt(decoded)
                            else:
                                text_input = decoded
                        st.success(f"**{uploaded_text.name}**（{len(text_input)} 文字）")
                    except Exception as exc:  # noqa: BLE001
                        st.error(f"読み込み失敗: {exc}")
                if text_input and text_input.strip():
                    transcript_text = text_input.strip()

    has_input = (audio_bytes is not None) or (transcript_text is not None)

    # ── Right: Status & submit ────────────────────────────────────────────────
    with col_status:
        with st.container(height=_vh(46), border=False):
            st.markdown('<p class="section-label">入力状況</p>', unsafe_allow_html=True)

            if audio_bytes is not None:
                st.markdown(
                    '<div class="status-panel">'
                    '<p class="status-ready">音声ファイル準備完了</p></div>',
                    unsafe_allow_html=True,
                )
                st.caption(f"ファイル: {filename}　サイズ: {len(audio_bytes)/1024:.1f} KB")
                use_batch = st.toggle(
                    "Batch Transcription",
                    value=False,
                    help="OFF: Fast（同期・高速）/ ON: Batch（非同期・大量処理向け）",
                )
                transcription_mode = "batch" if use_batch else "fast"
            elif transcript_text is not None:
                st.markdown(
                    '<div class="status-panel">'
                    '<p class="status-ready">テキスト準備完了</p></div>',
                    unsafe_allow_html=True,
                )
                st.caption(f"文字数: {len(transcript_text)}")
                transcription_mode = "fast"
            else:
                st.markdown(
                    '<div class="status-panel">'
                    '<p class="status-waiting">左のパネルで入力を選択してください</p></div>',
                    unsafe_allow_html=True,
                )
                transcription_mode = "fast"

            st.divider()
            if st.button(
                "議事録を生成",
                use_container_width=True,
                disabled=not has_input,
                type="primary",
            ):
                if transcript_text is not None:
                    st.session_state.input_mode = "transcript"
                    st.session_state.transcript_text = transcript_text
                    st.session_state.audio_bytes = None
                    st.session_state.audio_filename = None
                    st.session_state.audio_mime = None
                else:
                    st.session_state.input_mode = "audio"
                    st.session_state.transcript_text = None
                    st.session_state.audio_bytes = audio_bytes
                    st.session_state.audio_filename = filename
                    st.session_state.audio_mime = mime
                    st.session_state.transcription_mode = transcription_mode
                st.session_state.page = "processing"
                st.rerun()

    # ── History (lower half) ──────────────────────────────────────────────────
    _render_history_section()


def _render_history_section() -> None:
    """Render the saved-history section — always visible in the lower half."""
    st.markdown("---")
    col_title, col_reload = st.columns([5, 2])
    with col_title:
        st.markdown('<p class="section-label">過去の議事録</p>', unsafe_allow_html=True)
    with col_reload:
        if st.button("🔄 履歴リスト更新", key="reload_history", use_container_width=True):
            for k in list(st.session_state.keys()):
                if k == "_history_items" or k.startswith("_md_") or k.startswith("_in_"):
                    st.session_state.pop(k, None)
            st.rerun()

    items = st.session_state.get("_history_items")
    if items is None:
        items = api_list_history()
        st.session_state["_history_items"] = items

    if not items:
        st.info("まだ履歴はありません。議事録を生成すると自動で保存されます。")
        return

    _STEP_LABELS = {"step1": "音声解析", "step2": "スクリプト", "step3": "議事録", "step4": "用語補足"}

    with st.container(height=_vh(46), border=False):
        for it in items:
            jid = it.get("job_id", "")
            title = it.get("title") or "(無題)"
            created = it.get("created_at", "")
            try:
                created_disp = (
                    datetime.fromisoformat(created.replace("Z", "+00:00"))
                    .astimezone()
                    .strftime("%Y-%m-%d %H:%M")
                )
            except Exception:
                created_disp = created
            kind = it.get("input_kind", "")
            input_filename = it.get("input_filename", "")
            kind_label = "音声" if kind == "audio" else "テキスト"
            tmode = it.get("transcription_mode", "")
            tmode_label = ""
            if tmode == "batch":
                tmode_label = " ・ Batch"
            elif tmode == "fast":
                tmode_label = " ・ Fast"

            step_durs = it.get("step_durations") or {}
            total_dur = sum(step_durs.values()) if step_durs else None
            dur_total = f"合計 {total_dur:.0f}s" if total_dur is not None else ""
            # Build per-step breakdown string
            dur_parts = []
            for sk, sl in _STEP_LABELS.items():
                v = step_durs.get(sk)
                if v is not None:
                    dur_parts.append(f"{sl} {v:.1f}s")
            dur_breakdown = f"（{'／'.join(dur_parts)}）" if dur_parts else ""

            with st.container(border=True):
                is_timeout = bool(it.get("transcription_url"))
                st.markdown(
                    f"**{title}**  \n"
                    f"<span style='color:#605e5c;font-size:0.82rem;'>"
                    f"{created_disp} ・ {kind_label}（{input_filename}）{tmode_label}"
                    f" ・ <code>{jid[:8]}</code>"
                    f"{'  ・ <b style=\"color:#a4262c;\">タイムアウト</b>' if is_timeout else ''}"
                    f"</span>",
                    unsafe_allow_html=True,
                )
                if dur_total:
                    st.caption(f"{dur_total} {dur_breakdown}")

                if is_timeout:
                    # Timed-out job: show status check & resume buttons
                    tc1, tc2, tc3 = st.columns([1, 1, 1])
                    with tc1:
                        if st.button("状況確認", key=f"check_{jid}", use_container_width=True):
                            speech_status = api_check_transcription_status(jid)
                            if speech_status:
                                if speech_status == "Succeeded":
                                    st.success(f"文字起こし完了（{speech_status}）— 「再開」で議事録を生成できます")
                                elif speech_status in ("Running", "NotStarted"):
                                    st.info(f"文字起こし進行中（{speech_status}）")
                                elif speech_status == "Failed":
                                    st.error(f"文字起こし失敗（{speech_status}）")
                                else:
                                    st.warning(f"ステータス: {speech_status}")
                    with tc2:
                        if st.button("再開", key=f"resume_{jid}", use_container_width=True, type="primary"):
                            new_job_id = api_resume_transcription(jid)
                            if new_job_id:
                                st.session_state.job_id = new_job_id
                                st.session_state.input_mode = "audio"
                                st.session_state.page = "processing"
                                # Clear history cache
                                for k in list(st.session_state.keys()):
                                    if k == "_history_items" or k.startswith("_md_") or k.startswith("_in_"):
                                        st.session_state.pop(k, None)
                                st.rerun()
                    with tc3:
                        confirm_key = f"_del_confirm_{jid}"
                        if st.session_state.get(confirm_key):
                            if st.button("本当に削除", key=f"del2_{jid}", use_container_width=True):
                                if api_delete_history(jid):
                                    st.session_state.pop(confirm_key, None)
                                    st.session_state.pop("_history_items", None)
                                    st.rerun()
                        else:
                            if st.button("削除", key=f"del1_{jid}", use_container_width=True):
                                st.session_state[confirm_key] = True
                                st.rerun()
                else:
                    # Normal completed job
                    # Pre-fetch markdown + input lazily, cached per session.
                    md_key = f"_md_{jid}"
                    if md_key not in st.session_state:
                        meta = api_get_history(jid)
                        final = (meta or {}).get("result", {}).get("final_minutes") or {}
                        st.session_state[md_key] = final.get("markdown") or ""
                    in_key = f"_in_{jid}"
                    if in_key not in st.session_state:
                        payload = api_get_history_input(jid)
                        st.session_state[in_key] = payload  # tuple or None

                    md_data = st.session_state.get(md_key) or ""
                    in_payload = st.session_state.get(in_key)

                    c1, c2, c3, c4 = st.columns([1, 2, 2, 1])
                    with c1:
                        if st.button("開く", key=f"open_{jid}", use_container_width=True):
                            meta = api_get_history(jid)
                            if meta:
                                result = meta.get("result", {})
                                st.session_state.job_result = result
                                st.session_state.job_id = jid
                                st.session_state.input_mode = (
                                    "transcript" if kind == "transcript" else "audio"
                                )
                                st.session_state.page = "result"
                                st.rerun()
                    with c2:
                        st.download_button(
                            "議事録MD",
                            data=md_data.encode("utf-8") if md_data else b"",
                            file_name=f"minutes_{jid[:8]}.md",
                            mime="text/markdown",
                            key=f"dl_md_{jid}",
                            use_container_width=True,
                            disabled=not md_data,
                        )
                    with c3:
                        if in_payload:
                            data, fname, mime_ = in_payload
                            st.download_button(
                                "入力ファイル",
                                data=data,
                                file_name=fname,
                                mime=mime_,
                                key=f"dl_in_{jid}",
                                use_container_width=True,
                            )
                        else:
                            st.button(
                                "入力ファイル",
                                key=f"dl_in_disabled_{jid}",
                                use_container_width=True,
                                disabled=True,
                            )
                    with c4:
                        confirm_key = f"_del_confirm_{jid}"
                        if st.session_state.get(confirm_key):
                            if st.button(
                                "本当に削除",
                                key=f"del_yes_{jid}",
                                use_container_width=True,
                                type="primary",
                            ):
                                if api_delete_history(jid):
                                    for k in (md_key, in_key, confirm_key, "_history_items"):
                                        st.session_state.pop(k, None)
                                    st.rerun()
                        else:
                            if st.button(
                                "削除",
                                key=f"del_{jid}",
                                use_container_width=True,
                            ):
                                st.session_state[confirm_key] = True
                                st.rerun()


def page_processing() -> None:
    render_header(status="処理中...")

    _ALL_STEPS = [
        ("cu",      "", "音声解析（Speech Transcription）",
         "音声を文字起こしし構造化データを抽出"),
        ("script",  "", "スクリプト生成エージェント",
         "文字起こし結果を整理しスクリプト化"),
        ("minutes", "", "議事録作成エージェント",
         "スクリプトから正式な議事録を作成"),
        ("term",    "", "用語補足エージェント",
         "業界・社内用語を参照し議事録を補足"),
    ]
    input_mode = st.session_state.get("input_mode", "audio")
    if input_mode == "transcript":
        _STEPS = [s for s in _ALL_STEPS if s[0] != "cu"]
    else:
        _STEPS = _ALL_STEPS

    col_left, col_right = st.columns(2)

    # ── Left: step cards ──────────────────────────────────────────────────────
    with col_left:
        st.markdown('<p class="section-label">処理ステップ</p>', unsafe_allow_html=True)
        step_phs: dict[str, st.empty] = {}
        for key, icon, title, desc in _STEPS:
            ph = st.empty()
            step_phs[key] = ph
            ph.markdown(
                f'<div class="step-card pending">'
                f'<div><strong>{title}</strong><br>'
                f'<span style="font-size:.83rem;color:#a19f9d">{desc}</span></div>'
                f'<span style="margin-left:auto;font-size:.83rem;color:#a19f9d">待機中</span>'
                f"</div>",
                unsafe_allow_html=True,
            )
        msg_ph = st.empty()
        if input_mode == "transcript":
            msg_ph.info("文字起こしテキストを送信中...")
        else:
            msg_ph.info("音声ファイルを送信中...")

    # ── Right: agent detail (always visible) ──────────────────────────────────
    with col_right:
        st.markdown('<p class="section-label">エージェント詳細</p>', unsafe_allow_html=True)
        detail_ph = st.empty()

    def render_step(key: str, icon: str, title: str, desc: str,
                    state: str, status_text: str, error_detail: str = "") -> None:
        css = {"done": "done", "active": "active", "error": "done"}.get(state, "pending")
        if state == "error":
            css = "error"
        icon_prefix = {
            "done": "✔", "active": "▶", "error": "✖", "skipped": "–",
        }.get(state, "–")
        color = {"pending": "#a19f9d", "skipped": "#a19f9d"}.get(state, "inherit")
        detail_html = ""
        if error_detail:
            detail_html = (
                f'<div style="font-size:.8rem;color:#a4262c;margin-top:2px;'
                f'word-break:break-word">{error_detail}</div>'
            )
        step_phs[key].markdown(
            f'<div class="step-card {css}">'
            f'<div><strong>{title}</strong><br>'
            f'<span style="font-size:.83rem">{desc}</span>{detail_html}</div>'
            f'<span style="margin-left:auto;font-size:.83rem;color:{color};white-space:nowrap">'
            f"{icon_prefix} {status_text}</span>"
            f"</div>",
            unsafe_allow_html=True,
        )

    # ── Submit job ────────────────────────────────────────────────────────────
    job_id = st.session_state.job_id
    if job_id is None:
        try:
            if st.session_state.input_mode == "transcript":
                job_id = api_submit_transcript(st.session_state.transcript_text or "")
            else:
                job_id = api_submit(
                    st.session_state.audio_bytes,
                    st.session_state.audio_filename or "audio.wav",
                    st.session_state.audio_mime or "audio/wav",
                    transcription_mode=st.session_state.get("transcription_mode", "fast"),
                )
            st.session_state.job_id = job_id
        except Exception as exc:  # noqa: BLE001
            st.session_state.error_msg = f"送信エラー: {exc}"
            st.session_state.page = "error"
            st.rerun()
            return

    # ── Polling loop ──────────────────────────────────────────────────────────
    consecutive_errors = 0
    MAX_CONSEC_ERRORS = 5
    for _ in range(MAX_POLLS):
        try:
            result = api_poll(job_id)
            consecutive_errors = 0
        except Exception as exc:  # noqa: BLE001
            consecutive_errors += 1
            if consecutive_errors >= MAX_CONSEC_ERRORS:
                st.session_state.error_msg = (
                    f"ステータス取得エラーが連続しました: {exc}"
                )
                st.session_state.page = "error"
                st.rerun()
                return
            msg_ph.warning(
                f"⚠️ ステータス取得に失敗 ({consecutive_errors}/{MAX_CONSEC_ERRORS}回目) — 再試行します..."
            )
            time.sleep(POLL_INTERVAL)
            continue

        has_cu     = result.get("content_analysis") is not None
        has_script = result.get("script") is not None
        has_min    = result.get("minutes") is not None
        has_final  = result.get("final_minutes") is not None
        is_error   = result["status"] == "error"
        message    = result.get("message", "処理中...")
        step_durations = result.get("step_durations") or {}

        _DURATION_MAP = {"cu": "step1", "script": "step2", "minutes": "step3", "term": "step4"}

        failed_step: str | None = None
        if is_error:
            if input_mode == "transcript":
                if not has_script:
                    failed_step = "script"
                elif not has_min:
                    failed_step = "minutes"
                elif not has_final:
                    failed_step = "term"
            else:
                if not has_cu:
                    failed_step = "cu"
                elif not has_script:
                    failed_step = "script"
                elif not has_min:
                    failed_step = "minutes"
                elif not has_final:
                    failed_step = "term"

        error_reason = ""
        if is_error and message:
            if ":" in message:
                error_reason = message.split(":", 1)[1].strip()
            else:
                error_reason = message

        for key, icon, title, desc in _STEPS:
            if key == "cu":
                done, active = has_cu, not has_cu and not is_error
            elif key == "script":
                if input_mode == "transcript":
                    done, active = has_script, not has_script and not is_error
                else:
                    done, active = has_script, has_cu and not has_script and not is_error
            elif key == "minutes":
                done, active = has_min, has_script and not has_min and not is_error
            else:
                done, active = has_final, has_min and not has_final and not is_error

            if done:
                dur_key = _DURATION_MAP.get(key, "")
                dur_val = step_durations.get(dur_key)
                dur_text = f" ({dur_val:.1f}s)" if dur_val is not None else ""
                state, status_text, detail = "done", f"完了{dur_text}", ""
            elif is_error and key == failed_step:
                state, status_text = "error", "エラー"
                detail = error_reason
            elif is_error:
                state, status_text, detail = "skipped", "未到達", ""
            elif active:
                state, status_text, detail = "active", "処理中...", ""
            else:
                state, status_text, detail = "pending", "待機中", ""

            render_step(key, icon, title, desc, state, status_text, detail)

        # Update detail panel (always visible)
        with detail_ph.container():
            render_agent_detail(result, input_mode=input_mode)

        if result["status"] == "done":
            msg_ph.success("議事録が生成されました！")
            st.session_state.job_result = result
            st.session_state.page = "result"
            time.sleep(0.8)
            st.rerun()
            return

        if is_error:
            msg_ph.error(f"{message}")
            with col_left:
                if st.button("やり直す", use_container_width=True, type="primary"):
                    for k, v in _DEFAULTS.items():
                        st.session_state[k] = v
                    st.rerun()
            return

        msg_ph.info(f"{message}")
        time.sleep(POLL_INTERVAL)

    st.session_state.error_msg = "タイムアウト: 処理が完了しませんでした。"
    st.session_state.page = "error"
    st.rerun()


def page_result() -> None:
    render_header(status="完了")

    result = st.session_state.job_result or {}
    final  = result.get("final_minutes") or {}
    mins   = result.get("minutes") or {}
    script = result.get("script") or {}
    cu     = result.get("content_analysis") or {}

    markdown_text = final.get("markdown") or mins.get("raw_markdown", "")
    glossary      = final.get("glossary", [])

    col_left, col_right = st.columns(2)

    # ── Left: results ─────────────────────────────────────────────────────────
    with col_left:
        with st.container(height=_vh(46), border=False):
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    "Markdown",
                    data=markdown_text,
                    file_name=f"minutes_{_today()}.md",
                    mime="text/markdown",
                    use_container_width=True,
                )
            with col2:
                if st.button("新しい議事録", use_container_width=True):
                    for k, v in _DEFAULTS.items():
                        st.session_state[k] = v
                    st.rerun()

            t_min, t_script, t_transcript, t_glossary = st.tabs(
                ["議事録", "スクリプト", "文字起こし", "用語集"]
            )

            with t_min:
                if mins:
                    _render_structured_minutes(mins, glossary)
                elif markdown_text:
                    with st.container(border=True):
                        st.markdown(markdown_text)
                else:
                    st.info("議事録データがありません。")

            with t_script:
                s = script.get("script", "")
                if s:
                    st.text_area("会議スクリプト", value=s, height=180, disabled=True)
                    col_a, col_b = st.columns(2)
                    with col_a:
                        st.write("**参加者**")
                        for p in script.get("participants", []):
                            st.markdown(f"- {p}")
                    with col_b:
                        st.write("**議題**")
                        for a in script.get("agenda_items", []):
                            st.markdown(f"- {a}")
                else:
                    st.info("スクリプトデータがありません。")

            with t_transcript:
                tr = cu.get("raw_transcript", "")
                if tr:
                    st.text_area("生の文字起こし", value=tr, height=180, disabled=True)
                    meta_cols = st.columns(3)
                    with meta_cols[0]:
                        st.metric("話者数", len(cu.get("speakers", [])))
                    with meta_cols[1]:
                        dur = cu.get("duration_seconds")
                        st.metric("録音時間", f"{dur:.0f} 秒" if dur else "—")
                    with meta_cols[2]:
                        st.metric("言語", cu.get("language") or "—")
                    if cu.get("speakers"):
                        st.write("**話者:** " + "、".join(cu["speakers"]))
                    if cu.get("topics"):
                        st.write("**主なトピック:** " + "、".join(cu["topics"]))
                else:
                    st.info("文字起こしデータがありません。")

            with t_glossary:
                if glossary:
                    import pandas as pd
                    df = pd.DataFrame(glossary, columns=["term", "definition"])
                    df.columns = ["用語", "定義"]
                    st.dataframe(df, use_container_width=True, hide_index=True)
                else:
                    st.info("議事録中に専門用語は検出されませんでした。")

    # ── Right: agent detail (always visible) ──────────────────────────────────
    with col_right:
        render_agent_detail(
            result,
            input_mode=st.session_state.get("input_mode", "audio"),
            height=_vh(44),
        )

    # ── History (lower half) ──────────────────────────────────────────────────
    _render_history_section()


def page_error() -> None:
    render_header(status="エラー")
    st.error(f"エラー: {st.session_state.error_msg}")
    if st.button("やり直す", type="primary"):
        for k, v in _DEFAULTS.items():
            st.session_state[k] = v
        st.rerun()


# ── Helpers ───────────────────────────────────────────────────────────────────

def _today() -> str:
    from datetime import date
    return date.today().isoformat()


def _render_structured_minutes(mins: dict, glossary: list | None = None) -> None:
    """Render minutes in a Teams-style structure: 概要 / 議事 (collapsible topics) / フォローアップ タスク."""
    title = mins.get("title") or "会議議事録"
    date_str = mins.get("date") or ""
    participants = mins.get("participants") or []
    summary = (mins.get("summary") or "").strip()
    topics = mins.get("topics") or []
    follow_up = mins.get("follow_up_tasks") or mins.get("action_items") or []

    with st.container(border=True):
        st.markdown(f"### {title}")
        meta_parts = []
        if date_str:
            meta_parts.append(f"**日時：** {date_str}")
        if participants:
            meta_parts.append(f"**参加者：** {'、'.join(participants)}")
        if meta_parts:
            st.markdown("　".join(meta_parts))

        # ── 概要 ──
        st.markdown("#### 概要")
        if summary:
            st.markdown(summary)
        else:
            st.caption("概要は生成されませんでした。")

        st.markdown("---")

        # ── 議事 (collapsible topics) ──
        col_h, col_btn = st.columns([6, 1])
        with col_h:
            st.markdown("#### 議事")
        with col_btn:
            expand_all = st.toggle("すべて展開", key="topics_expand_all", value=False)

        if topics:
            for i, t in enumerate(topics, 1):
                t_title = t.get("title") or f"トピック {i}"
                t_summary = (t.get("summary") or "").strip()
                t_details = t.get("details") or []
                with st.expander(f"**{i}. {t_title}**", expanded=expand_all):
                    if t_summary:
                        st.markdown(
                            f"<div class='topic-summary'>{t_summary}</div>",
                            unsafe_allow_html=True,
                        )
                    for d in t_details:
                        st.markdown(f"- {d}")
        else:
            # Fallback: render the raw markdown body for legacy minutes without topics.
            raw = (mins.get("raw_markdown") or "").strip()
            if raw:
                st.markdown(raw)
            else:
                st.caption("議事は生成されませんでした。")

        st.markdown("---")

        # ── フォローアップ タスク ──
        st.markdown("#### フォローアップ タスク")
        if follow_up:
            for t in follow_up:
                task = t.get("task", "")
                owner = t.get("owner")
                due = t.get("due")
                meta = []
                if owner:
                    meta.append(f"担当: {owner}")
                if due:
                    meta.append(f"期限: {due}")
                meta_text = f"（{' / '.join(meta)}）" if meta else ""
                st.markdown(f"- **{task}** {meta_text}")
        else:
            st.caption("フォローアップ タスクはありません。")

        # ── 用語集（あれば） ──
        if glossary:
            st.markdown("---")
            st.markdown("#### 用語集")
            import pandas as pd
            df = pd.DataFrame(glossary, columns=["term", "definition"])
            df.columns = ["用語", "定義"]
            st.dataframe(df, use_container_width=True, hide_index=True)


# ── Router ────────────────────────────────────────────────────────────────────

def main() -> None:
    page = st.session_state.page
    if page == "input":
        page_input()
    elif page == "processing":
        page_processing()
    elif page == "result":
        page_result()
    elif page == "error":
        page_error()

    # ── Footer ────────────────────────────────────────────────────────────────
    st.markdown(
        "<p style='text-align:center;color:#a19f9d;font-size:.78rem;margin-top:8px'>"
        "Meeting Minutes Agent — Azure AI × Azure OpenAI × Container Apps"
        "</p>",
        unsafe_allow_html=True,
    )


if __name__ == "__main__":
    main()
